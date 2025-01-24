
# # # components/cv_analyzer.py
# # import streamlit as st
# # from PyPDF2 import PdfReader
# # import docx
# # from openai import OpenAI


# # button_css = """
# # <style>
# # div.stButton > button {
# #     display: block;
# #     margin: 0 auto;
# # }
# # </style>
# # """


# # def read_pdf(file):
# #     pdf = PdfReader(file)
# #     text = ""
# #     for page in pdf.pages:
# #         text += page.extract_text()
# #     return text

# # def read_docx(file):
# #     doc = docx.Document(file)
# #     text = ""
# #     for para in doc.paragraphs:
# #         text += para.text + "\n"
# #     return text

# # def analyze_cv(text, api_key):
# #     client = OpenAI(api_key=api_key)
# #     response = client.chat.completions.create(
# #         model="gpt-4",
# #         messages=[
# #             {"role": "system", "content": "Analyze this CV and provide insights."},
# #             {"role": "user", "content": text}
# #         ]
# #     )
# #     return response.choices[0].message.content

# # def get_cv_improvements(text, api_key):
# #     client = OpenAI(api_key=api_key)
# #     response = client.chat.completions.create(
# #         model="gpt-4",
# #         messages=[
# #             {"role": "system", "content": "Suggest improvements for this CV."},
# #             {"role": "user", "content": text}
# #         ]
# #     )
# #     return response.choices[0].message.content

# # def render_cv_section():
# #     st.header("CV Analysis")
# #     uploaded_file = st.file_uploader("Upload your CV", type=['pdf', 'docx'])
    
# #     if uploaded_file:
# #         text = ""
# #         if uploaded_file.type == "application/pdf":
# #             text = read_pdf(uploaded_file)
# #         else:
# #             text = read_docx(uploaded_file)
# #         # print(f"text ==  {text}")

# #         st.markdown(button_css, unsafe_allow_html=True)

# #         if st.button("Get cv raw content"):
# #             st.text_area("CV Content", text, height=200)
         
        
# #     if st.button("Analyze CV"):
# #         api_key = st.session_state.get('openai_api_key')
# #         if not api_key:
# #             st.error("Please enter OpenAI API key in sidebar")
# #             return
            
# #         analysis = analyze_cv(text, api_key)
# #         improvements = get_cv_improvements(text, api_key)
        
# #         st.subheader("Analysis")
# #         st.write(analysis)



# #     if st.button("Get improvement from OpenAI"):
# #         api_key = st.session_state.get('openai_api_key')
# #         if not api_key:
# #             st.error("Please enter OpenAI API key in sidebar")
# #             return
            
# #         improvements = get_cv_improvements(text, api_key)
# #         st.subheader("Suggested Improvements")
# #         st.write(improvements)





# import streamlit as st
# from PyPDF2 import PdfReader
# import docx
# from openai import OpenAI

# button_css = """
# <style>
# div.stButton > button {
#     display: block;
#     margin: 0 auto;
# }
# </style>
# """

# def read_pdf(file):
#     pdf = PdfReader(file)
#     return "".join(page.extract_text() for page in pdf.pages)

# def read_docx(file):
#     doc = docx.Document(file)
#     return "\n".join(para.text for para in doc.paragraphs)

# def analyze_cv(text, api_key):
#     client = OpenAI(api_key=api_key)
#     response = client.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "Analyze this CV and provide insights."},
#             {"role": "user", "content": text}
#         ]
#     )
#     return response.choices[0].message.content

# def get_cv_improvements(text, api_key):
#     client = OpenAI(api_key=api_key)
#     response = client.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "Suggest improvements for this CV."},
#             {"role": "user", "content": text}
#         ]
#     )
#     return response.choices[0].message.content

# def render_cv_section():
#     st.header("CV Analysis")
    
#     if 'cv_text' not in st.session_state:
#         st.session_state.cv_text = ""
#     if 'analysis_result' not in st.session_state:
#         st.session_state.analysis_result = ""
#     if 'improvement_result' not in st.session_state:
#         st.session_state.improvement_result = ""
#     if 'show_raw_content' not in st.session_state:
#         st.session_state.show_raw_content = False

#     uploaded_file = st.file_uploader("Upload your CV", type=['pdf', 'docx'])
    
#     if uploaded_file:
#         if uploaded_file.type == "application/pdf":
#             st.session_state.cv_text = read_pdf(uploaded_file)
#         else:
#             st.session_state.cv_text = read_docx(uploaded_file)
    
#     st.markdown(button_css, unsafe_allow_html=True)
    
#     if st.button("Get cv raw content") and st.session_state.cv_text:
#         st.session_state.show_raw_content = True
        
#     if st.session_state.show_raw_content and st.session_state.cv_text:
#         st.text_area("CV Content", st.session_state.cv_text, height=200)

#     analyze_button = st.button("Analyze CV")
#     if analyze_button and st.session_state.cv_text:
#         api_key = st.session_state.get('openai_api_key')
#         if not api_key:
#             st.error("Please enter OpenAI API key in sidebar")
#         else:    
#             st.session_state.analysis_result = analyze_cv(st.session_state.cv_text, api_key)
    
#     if st.session_state.analysis_result:
#         st.subheader("Analysis")
#         st.write(st.session_state.analysis_result)

#     improve_button = st.button("Get improvement from OpenAI")
#     if improve_button and st.session_state.cv_text:
#         api_key = st.session_state.get('openai_api_key')
#         if not api_key:
#             st.error("Please enter OpenAI API key in sidebar")
#         else:
#             st.session_state.improvement_result = get_cv_improvements(st.session_state.cv_text, api_key)
    
#     if st.session_state.improvement_result:
#         st.subheader("Suggested Improvements") 
#         st.write(st.session_state.improvement_result)














# -------------------------------------versioon2---------------------------

import streamlit as st
from PyPDF2 import PdfReader
import docx
from openai import OpenAI
from groq import Groq

button_css = """
<style>
div.stButton > button {
    display: block;
    margin: 0 auto;
}
</style>
"""

def read_pdf(file):
    pdf = PdfReader(file)
    return "".join(page.extract_text() for page in pdf.pages)

def read_docx(file):
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

def analyze_with_groq(text, api_key):
    client = Groq(api_key=api_key)
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "Analyze this CV and provide insights."},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content

def analyze_cv(text, api_key, api_type='openai'):
    if api_type == 'openai':
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Analyze this CV and provide insights."},
                {"role": "user", "content": text}
            ]
        )
        return response.choices[0].message.content
    else:
        return analyze_with_groq(text, api_key)

def get_cv_improvements(text, api_key, api_type='openai'):
    if api_type == 'openai':
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Suggest improvements for this CV."},
                {"role": "user", "content": text}
            ]
        )
        return response.choices[0].message.content
    else:
        return analyze_with_groq(text, api_key)

def render_cv_section():
    st.header("CV Analysis")
    
    if 'cv_text' not in st.session_state:
        st.session_state.cv_text = ""
    if 'analysis_result' not in st.session_state:
        st.session_state.analysis_result = ""
    if 'improvement_result' not in st.session_state:
        st.session_state.improvement_result = ""
    if 'show_raw_content' not in st.session_state:
        st.session_state.show_raw_content = False

    uploaded_file = st.file_uploader("Upload your CV", type=['pdf', 'docx'])
    
    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            st.session_state.cv_text = read_pdf(uploaded_file)
        else:
            st.session_state.cv_text = read_docx(uploaded_file)
    
    st.markdown(button_css, unsafe_allow_html=True)
    
    if st.button("Get cv raw content") and st.session_state.cv_text:
        st.session_state.show_raw_content = True
        
    if st.session_state.show_raw_content and st.session_state.cv_text:
        st.text_area("CV Content", st.session_state.cv_text, height=200)

    analyze_button = st.button("Analyze CV")
    if analyze_button and st.session_state.cv_text:
        api_key = st.session_state.get('api_key')
        if not api_key:
            st.error(f"Please enter and activate an API key in sidebar")
        else:    
            st.session_state.analysis_result = analyze_cv(
                st.session_state.cv_text, 
                api_key, 
                st.session_state.active_api
            )
    
    if st.session_state.analysis_result:
        st.subheader("Analysis")
        st.write(st.session_state.analysis_result)

    improve_button = st.button("Get improvement from OpenAI")
    if improve_button and st.session_state.cv_text:
        api_key = st.session_state.get('api_key')
        if not api_key:
            st.error("Please enter and activate an API key in sidebar")
        else:
            st.session_state.improvement_result = get_cv_improvements(
                st.session_state.cv_text,
                api_key,
                st.session_state.active_api
            )
    
    if st.session_state.improvement_result:
        st.subheader("Suggested Improvements") 
        st.write(st.session_state.improvement_result)